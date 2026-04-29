"""Genera el reporte de pruebas en formato .docx para el plan de pruebas."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_paragraph(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def make_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1F4E78")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_cm:
        for row in table.rows:
            for i, width in enumerate(col_widths_cm):
                row.cells[i].width = Cm(width)
    return table


def main():
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Margenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ---------- Portada / Titulo ----------
    title = doc.add_heading("Reporte de Ejecucion de Pruebas", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Sistema de Registro de Ventas Cahuamanta \"El Chuta\"")
    run.bold = True
    run.font.size = Pt(13)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Casos de uso del rol Dueno")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ---------- Datos generales ----------
    add_heading(doc, "1. Datos generales", level=1)
    info = [
        ("Modulo bajo prueba", "Negocio (capa de logica de negocio)"),
        ("Framework de pruebas", "JUnit 5.10.3"),
        ("Herramienta de ejecucion", "Apache Maven 3.9.15 - Surefire 3.5.4"),
        ("JDK", "21"),
        ("Fecha de ejecucion", "2026-04-29"),
        ("Total de pruebas", "41"),
        ("Resultado global", "41 OK / 0 Fallos / 0 Errores / 0 Omitidas"),
        ("Tiempo total", "6.581 s"),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].text = ""
        run = t.rows[i].cells[0].paragraphs[0].add_run(k)
        run.bold = True
        run.font.size = Pt(10)
        shade_cell(t.rows[i].cells[0], "DEEAF6")
        t.rows[i].cells[1].text = ""
        t.rows[i].cells[1].paragraphs[0].add_run(v).font.size = Pt(10)

    doc.add_paragraph()

    # ---------- CU01 ----------
    add_heading(doc, "2. CU01 - Gestionar Usuarios", level=1)
    add_paragraph(doc, "Clase de prueba: Pruebaa.UsuarioBOTest", bold=True, size=10)
    add_paragraph(
        doc,
        "Metodos de fachada cubiertos: crearUsuario, actualizarUsuario, "
        "eliminarUsuario, obtenerUsuario, obtenerUsuarios.",
        size=10,
    )

    cu01 = [
        ("CP-U-01", "crearUsuarioPersisteCajeroValido",
         "nombre=\"juan\", contrasenia=\"123\", rol=\"CAJERO\"",
         "Se persiste 1 usuario con rol CAJERO", "OK"),
        ("CP-U-02", "crearUsuarioNormalizaRolDuenoAAdmin",
         "nombre=\"ana\", rol=\"DUENO\"",
         "Se persiste con rol normalizado ADMIN", "OK"),
        ("CP-U-03", "crearUsuarioConNombreVacioLanzaError",
         "nombre=\"\", contrasenia=\"x\"",
         "NegocioException: El nombre de usuario es obligatorio.", "OK"),
        ("CP-U-04", "crearUsuarioConNombreEnBlancoLanzaError",
         "nombre=\"   \"",
         "NegocioException: El nombre de usuario es obligatorio.", "OK"),
        ("CP-U-05", "crearUsuarioConContraseniaVaciaLanzaError",
         "nombre=\"luis\", contrasenia=\"\"",
         "NegocioException: La contrasenia es obligatoria.", "OK"),
        ("CP-U-06", "crearUsuarioConRolInvalidoLanzaError",
         "rol=\"GERENTE\"",
         "NegocioException: Rol no valido.", "OK"),
        ("CP-U-07", "crearUsuarioConNombreDuplicadoLanzaError",
         "Existe \"Juan\"; se crea \"JUAN\"",
         "NegocioException: Ya existe un usuario con ese nombre.", "OK"),
        ("CP-U-08", "crearUsuarioConDTONuloLanzaError",
         "usuarioDTO = null",
         "NegocioException: Datos de usuario no enviados.", "OK"),
        ("CP-U-09", "obtenerUsuarioPorIdExistenteDevuelveDTO",
         "id = 1",
         "UsuarioDTO con datos correctos y rol normalizado", "OK"),
        ("CP-U-10", "obtenerUsuarioPorIdInexistenteDevuelveNull",
         "id = 9999",
         "Devuelve null sin lanzar excepcion", "OK"),
        ("CP-U-11", "obtenerTodosLosUsuariosDevuelveListaCompleta",
         "3 usuarios en BD",
         "Lista con 3 elementos y roles normalizados", "OK"),
        ("CP-U-12", "obtenerTodosLosUsuariosDevuelveListaVaciaCuandoNoHayRegistros",
         "Tabla vacia",
         "Lista vacia no nula", "OK"),
        ("CP-U-13", "actualizarUsuarioPersisteCambiosValidos",
         "id=1, nombre=\"juanp\"",
         "DAO recibe la actualizacion con los nuevos datos", "OK"),
        ("CP-U-14", "actualizarUsuarioSinIdLanzaError",
         "idUsuario = null",
         "NegocioException: Usuario no valido para actualizar.", "OK"),
        ("CP-U-15", "actualizarUsuarioConNombreTomadoPorOtroLanzaError",
         "id=2 quiere nombre \"juan\" (de id=1)",
         "NegocioException: Ya existe un usuario con ese nombre.", "OK"),
        ("CP-U-16", "actualizarUsuarioManteniendoSuMismoNombreEsPermitido",
         "id=1 conserva su nombre",
         "No lanza excepcion", "OK"),
        ("CP-U-17", "eliminarUsuarioExistenteInvocaDAO",
         "id = 2",
         "DAO recibe llamada eliminar(2)", "OK"),
        ("CP-U-18", "eliminarUsuarioInexistenteNoLanzaError",
         "id = 9999",
         "No lanza excepcion", "OK"),
    ]

    make_table(
        doc,
        ["ID", "Metodo de prueba", "Entrada", "Resultado esperado", "Resultado"],
        cu01,
        col_widths_cm=[1.6, 5.5, 4.0, 5.0, 1.4],
    )
    add_paragraph(doc, "Resumen CU01: 18 ejecutados / 18 OK / 0 fallos.", bold=True, size=10)

    doc.add_paragraph()

    # ---------- CU02 ----------
    add_heading(doc, "3. CU02 - Consultar Historial de Ventas", level=1)
    add_paragraph(doc, "Clase de prueba: Pruebaa.VentaBOHistorialTest", bold=True, size=10)
    add_paragraph(
        doc,
        "Metodos de fachada cubiertos: obtenerVentasDelDia, obtenerVentasPorFecha.",
        size=10,
    )

    cu02 = [
        ("CP-H-01", "obtenerVentasPorFechaDelDiaActualDevuelveVentas",
         "fecha = hoy; 3 ventas",
         "Lista con 3 ventas", "OK"),
        ("CP-H-02", "obtenerVentasPorFechaDelDiaSinRegistrosDevuelveListaVacia",
         "fecha = hoy; 0 ventas",
         "Lista vacia no nula", "OK"),
        ("CP-H-03", "obtenerVentasPorFechaPasadaDevuelveVentasDelDia",
         "fecha=2026-04-20; 5 ventas",
         "Lista con 5 ventas en orden", "OK"),
        ("CP-H-04", "obtenerVentasPorFechaFuturaDevuelveListaVacia",
         "fecha=2030-01-01",
         "Lista vacia", "OK"),
        ("CP-H-05", "obtenerVentasPorFechaNullUsaFechaActual",
         "fecha = null",
         "Consulta rango del dia actual [hoy 00:00, manana 00:00)", "OK"),
        ("CP-H-06", "obtenerVentasPorFechaConsultaRangoDelDiaCompleto",
         "fecha=2026-04-27",
         "Rango DAO: [2026-04-27 00:00, 2026-04-28 00:00)", "OK"),
    ]
    make_table(
        doc,
        ["ID", "Metodo de prueba", "Entrada", "Resultado esperado", "Resultado"],
        cu02,
        col_widths_cm=[1.6, 5.5, 4.0, 5.0, 1.4],
    )
    add_paragraph(doc, "Resumen CU02: 6 ejecutados / 6 OK / 0 fallos.", bold=True, size=10)

    doc.add_paragraph()

    # ---------- CU03 ----------
    add_heading(doc, "4. CU03 - Generar Reporte de Ventas", level=1)
    add_paragraph(doc, "Clase de prueba: Pruebaa.VentaBOReporteTest", bold=True, size=10)
    add_paragraph(
        doc,
        "Metodos de fachada cubiertos: obtenerReporteVentas, exportarReporteVentas.",
        size=10,
    )

    cu03 = [
        ("CP-R-01", "obtenerReporteVentasCalculaTotalesYPromedio",
         "2 ventas: 100.00 + 250.50",
         "cantidad=2, total=350.50, promedio=175.25", "OK"),
        ("CP-R-02", "obtenerReporteVentasSinResultadosDevuelveTotalesEnCero",
         "0 ventas en rango valido",
         "cantidad=0, total=0, promedio=0", "OK"),
        ("CP-R-03", "obtenerReporteVentasMismoDiaConsultaRangoCompleto",
         "inicio=fin=2026-04-29",
         "Rango DAO: [2026-04-29 00:00, 2026-04-30 00:00)", "OK"),
        ("CP-R-04", "obtenerReporteVentasConFechaInicioNullLanzaError",
         "inicio=null, fin=2026-04-29",
         "NegocioException: Debe seleccionar una fecha inicial y una fecha final.", "OK"),
        ("CP-R-05", "obtenerReporteVentasConFechaFinNullLanzaError",
         "inicio=2026-04-29, fin=null",
         "NegocioException: Debe seleccionar una fecha inicial y una fecha final.", "OK"),
        ("CP-R-06", "obtenerReporteVentasConRangoInvertidoLanzaError",
         "inicio=2026-04-29, fin=2026-04-01",
         "NegocioException: La fecha final no puede ser menor que la fecha inicial.", "OK"),
        ("CP-R-07", "obtenerReporteVentasCalculaPromedioRedondeadoHalfUp",
         "3 ventas: 100, 200, 300",
         "total=600.00, promedio=200.00", "OK"),
        ("CP-R-08", "obtenerReporteVentasRedondeaPromedioConDivisionNoExacta",
         "3 ventas: 100, 100, 100.01",
         "total=300.01, promedio=100.00 (HALF_UP, 2 dec.)", "OK"),
        ("CP-R-09", "obtenerReporteVentasIgnoraVentasConTotalNull",
         "1 venta total=null + 1 con total=100",
         "cantidad=2, total=100.00 (sin NPE)", "OK"),
        ("CP-R-10", "exportarReporteVentasGeneraPdfCuandoHayDatos",
         "rango con 1 venta, ruta valida",
         "PDF creado y con tamano > 0", "OK"),
        ("CP-R-11", "exportarReporteVentasSinDatosLanzaError",
         "rango sin ventas",
         "NegocioException: No hay registros disponibles en el periodo seleccionado.", "OK"),
        ("CP-R-12", "exportarReporteVentasConRutaNullLanzaError",
         "ruta = null",
         "NegocioException: La ruta del archivo no es valida.", "OK"),
        ("CP-R-13", "exportarReporteVentasConRutaEnBlancoLanzaError",
         "ruta = \"   \"",
         "NegocioException: La ruta del archivo no es valida.", "OK"),
        ("CP-R-14", "exportarReporteVentasCreaDirectoriosFaltantes",
         "ruta con carpetas inexistentes",
         "Crea las carpetas y genera el PDF", "OK"),
        ("CP-R-16", "exportarReporteVentasConCaracteresEspecialesNoFalla",
         "usuario=\"Jose Nunez\"",
         "PDF generado sin error (texto normalizado)", "OK"),
        ("CP-R-17", "exportarReporteVentasGeneraMultiplesPaginasConMuchasVentas",
         "60 ventas",
         "PDF generado con tamano > 1 KB (varias paginas)", "OK"),
        ("CP-R-18", "exportarReporteVentasConFolioNullNoLanzaError",
         "venta con folio=null",
         "PDF generado, sin NPE (muestra #<id>)", "OK"),
    ]
    make_table(
        doc,
        ["ID", "Metodo de prueba", "Entrada", "Resultado esperado", "Resultado"],
        cu03,
        col_widths_cm=[1.6, 5.5, 4.0, 5.0, 1.4],
    )
    add_paragraph(doc, "Resumen CU03: 17 ejecutados / 17 OK / 0 fallos.", bold=True, size=10)

    doc.add_paragraph()

    # ---------- Matriz de trazabilidad ----------
    add_heading(doc, "5. Matriz de trazabilidad", level=1)
    matriz = [
        ("CU01 - Gestionar usuarios", "CP-U-01 a CP-U-18", "18", "18", "0"),
        ("CU02 - Consultar historial", "CP-H-01 a CP-H-06", "6", "6", "0"),
        ("CU03 - Generar reporte", "CP-R-01 a CP-R-18", "17", "17", "0"),
        ("TOTAL", "-", "41", "41", "0"),
    ]
    make_table(
        doc,
        ["Caso de uso", "Casos de prueba", "Total", "OK", "Fallos"],
        matriz,
        col_widths_cm=[5.5, 5.5, 2.0, 2.0, 2.0],
    )

    doc.add_paragraph()

    # ---------- Bitacora ----------
    add_heading(doc, "6. Bitacora de defectos", level=1)
    add_paragraph(doc, "No se identificaron defectos durante la ejecucion de las pruebas.", size=11)

    # ---------- Evidencia ----------
    add_heading(doc, "7. Evidencia de ejecucion (consola Maven)", level=1)
    evidencia = (
        "-------------------------------------------------------\n"
        " T E S T S\n"
        "-------------------------------------------------------\n"
        "Running Pruebaa.UsuarioBOTest\n"
        "Tests run: 18, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.047 s\n"
        "Running Pruebaa.VentaBOHistorialTest\n"
        "Tests run: 6, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.016 s\n"
        "Running Pruebaa.VentaBOReporteTest\n"
        "Tests run: 17, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.458 s\n\n"
        "Results:\n"
        "Tests run: 41, Failures: 0, Errors: 0, Skipped: 0\n\n"
        "BUILD SUCCESS\n"
        "Total time:  6.581 s\n"
        "Finished at: 2026-04-29T12:31:34"
    )
    p = doc.add_paragraph()
    run = p.add_run(evidencia)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_paragraph()

    # ---------- Conclusiones ----------
    add_heading(doc, "8. Conclusiones", level=1)
    conclusiones = [
        "Las 41 pruebas unitarias ejecutadas para los tres casos de uso del rol Dueno "
        "pasaron sin fallos.",
        "Se cubren validaciones de entrada (nombres y contrasenias vacias, roles "
        "invalidos, fechas invalidas), reglas de negocio (duplicidad de nombres, "
        "normalizacion de roles, calculo de totales y promedio), rutas alternas "
        "(datos nulos, listas vacias) y exportacion a PDF (rutas inexistentes, "
        "caracteres especiales, paginacion).",
        "No se identificaron defectos. Se considera aceptado el modulo de Negocio "
        "para los casos de uso del rol Dueno.",
    ]
    for texto in conclusiones:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(texto).font.size = Pt(11)

    salida = r"C:\Users\chris\OneDrive\Documents\GitHub\Sistema-de-Registro-de-Ventas-Cahuamanta-El-Chuta-\Reporte_Pruebas_Dueno.docx"
    doc.save(salida)
    print(f"Documento generado: {salida}")


if __name__ == "__main__":
    main()
